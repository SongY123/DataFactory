from __future__ import annotations

import csv
import io
import mimetypes
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from charset_normalizer import from_bytes as detect_from_bytes
except Exception:  # pragma: no cover
    detect_from_bytes = None

from utils.config_loader import get_config
from utils.logger import logger


class FileParseError(ValueError):
    """Raised when uploaded file parsing fails with a user-safe message."""


@dataclass
class ParsedFileContent:
    file_name: str
    file_type: str
    mime_type: str
    size_bytes: int
    parser: str
    text_summary: str
    table_structures: list[dict[str, Any]]
    sample_content: list[str]
    warnings: list[str]
    metadata: dict[str, Any]

    def to_prompt_payload(self) -> dict[str, Any]:
        return {
            "file_name": self.file_name,
            "file_type": self.file_type,
            "mime_type": self.mime_type,
            "size_bytes": self.size_bytes,
            "parser": self.parser,
            "text_summary": self.text_summary,
            "table_structures": self.table_structures,
            "sample_content": self.sample_content,
            "warnings": self.warnings,
            "metadata": self.metadata,
        }


class AgentFileParser:
    _TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".jsonl"}
    _DOC_MIME_TYPES = {"application/msword"}

    def __init__(self) -> None:
        self._max_parse_seconds = int(self._get_config("agent.file.max_parse_seconds", 25) or 25)

        self._max_text_chars = int(self._get_config("agent.file.max_text_chars", 12000) or 12000)
        self._sample_rows = int(self._get_config("agent.file.sample_rows", 6) or 6)

        self._pdf_max_pages = int(self._get_config("agent.file.pdf_max_pages", 20) or 20)
        self._docx_max_paragraphs = int(self._get_config("agent.file.docx_max_paragraphs", 120) or 120)
        self._docx_max_tables = int(self._get_config("agent.file.docx_max_tables", 12) or 12)
        self._xlsx_max_sheets = int(self._get_config("agent.file.xlsx_max_sheets", 12) or 12)

        self._started_at = 0.0

    @staticmethod
    def _get_config(key: str, default: Any) -> Any:
        try:
            return get_config(key, default)
        except Exception:
            return default

    def _start_timer(self) -> None:
        self._started_at = time.monotonic()

    def _check_timeout(self) -> None:
        if (time.monotonic() - self._started_at) > self._max_parse_seconds:
            raise FileParseError(f"File parsing timed out (> {self._max_parse_seconds}s).")

    @staticmethod
    def _normalize_text(text: str) -> str:
        return "\n".join(line.rstrip() for line in str(text or "").replace("\x00", "").splitlines())

    @staticmethod
    def _truncate_text(text: str, limit: int) -> str:
        value = str(text or "")
        if len(value) <= limit:
            return value
        return value[:limit] + "\n...[truncated]"

    @staticmethod
    def _looks_like_binary(raw: bytes) -> bool:
        sample = raw[:4096]
        if not sample:
            return False
        if b"\x00" in sample:
            return True
        control = sum(1 for b in sample if b < 9 or (13 < b < 32))
        ratio = control / len(sample)
        return ratio > 0.30

    def _decode_text(self, raw: bytes) -> tuple[str, str]:
        if self._looks_like_binary(raw):
            raise FileParseError("File looks binary and cannot be decoded as plain text.")

        if detect_from_bytes is not None:
            best = detect_from_bytes(raw).best()
            if best is not None:
                encoding = str(best.encoding or "unknown")
                return self._normalize_text(str(best)), encoding

        for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
            try:
                return self._normalize_text(raw.decode(enc)), enc
            except UnicodeDecodeError:
                continue

        raise FileParseError("Text encoding detection failed. Please convert file to UTF-8.")

    @staticmethod
    def _guess_mime(file_name: str, content_type: str | None) -> str:
        provided = str(content_type or "").strip().lower()
        if provided:
            return provided
        guessed, _ = mimetypes.guess_type(file_name)
        return str(guessed or "application/octet-stream").lower()

    @staticmethod
    def _detect_kind(file_name: str, mime_type: str) -> str:
        ext = Path(file_name).suffix.lower()
        mime = mime_type.lower()

        if ext == ".doc" or mime in AgentFileParser._DOC_MIME_TYPES:
            return "doc"
        if ext == ".pdf" or mime == "application/pdf":
            return "pdf"
        if ext == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return "docx"
        if ext in {".xlsx", ".xlsm"} or mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return "xlsx"
        if ext in AgentFileParser._TEXT_EXTENSIONS:
            return "text"
        if mime.startswith("text/"):
            return "text"
        return "unsupported"

    @staticmethod
    def _stringify_cell(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, float):
            return f"{value:.6g}"
        return str(value).strip()

    def _parse_text(self, *, file_name: str, mime_type: str, raw: bytes) -> ParsedFileContent:
        text, encoding = self._decode_text(raw)
        ext = Path(file_name).suffix.lower()

        table_structures: list[dict[str, Any]] = []
        sample_content: list[str] = []
        metadata: dict[str, Any] = {
            "encoding": encoding,
            "line_count": text.count("\n") + 1 if text else 0,
        }

        if ext == ".csv" or "csv" in mime_type:
            reader = csv.reader(io.StringIO(text))
            rows: list[list[str]] = []
            for idx, row in enumerate(reader):
                self._check_timeout()
                rows.append([self._stringify_cell(v) for v in row])
                if idx >= self._sample_rows:
                    break

            header = rows[0] if rows else []
            sample_rows = rows[1 : self._sample_rows + 1] if len(rows) > 1 else []
            table_structures.append(
                {
                    "name": "csv_table",
                    "columns": header[:50],
                    "sample_rows": sample_rows,
                }
            )
            sample_content = [", ".join(row[:12]) for row in sample_rows[:3] if any(row)]
            metadata["column_count"] = len(header)

        if not sample_content:
            sample_content = [line.strip() for line in text.splitlines() if line.strip()][:6]

        summary = self._truncate_text(text, self._max_text_chars)
        return ParsedFileContent(
            file_name=Path(file_name).name,
            file_type=ext.lstrip(".") or "text",
            mime_type=mime_type,
            size_bytes=len(raw),
            parser="text",
            text_summary=summary,
            table_structures=table_structures,
            sample_content=sample_content,
            warnings=[],
            metadata=metadata,
        )

    def _parse_pdf(self, *, file_name: str, mime_type: str, raw: bytes) -> ParsedFileContent:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            logger.warning("PDF dependency missing: %s", exc)
            raise FileParseError("PDF parser is unavailable on server. Please contact administrator.") from exc

        try:
            reader = PdfReader(io.BytesIO(raw))
        except Exception as exc:
            logger.warning("Failed to load PDF file %s: %s", file_name, exc)
            raise FileParseError("PDF parse failed. Please confirm file is valid or export it again.") from exc

        texts: list[str] = []
        for idx, page in enumerate(reader.pages):
            self._check_timeout()
            if idx >= self._pdf_max_pages:
                break
            text = self._normalize_text(page.extract_text() or "")
            if text.strip():
                texts.append(text)

        merged = "\n\n".join(texts).strip()
        if not merged:
            raise FileParseError("No readable text extracted from PDF. If scanned, run OCR first.")

        return ParsedFileContent(
            file_name=Path(file_name).name,
            file_type="pdf",
            mime_type=mime_type,
            size_bytes=len(raw),
            parser="pypdf",
            text_summary=self._truncate_text(merged, self._max_text_chars),
            table_structures=[],
            sample_content=[line.strip() for line in merged.splitlines() if line.strip()][:8],
            warnings=[],
            metadata={
                "page_count": len(reader.pages),
                "parsed_pages": min(len(reader.pages), self._pdf_max_pages),
            },
        )

    def _parse_docx(self, *, file_name: str, mime_type: str, raw: bytes) -> ParsedFileContent:
        try:
            from docx import Document
        except Exception as exc:
            logger.warning("DOCX dependency missing: %s", exc)
            raise FileParseError("DOCX parser is unavailable on server. Please contact administrator.") from exc

        try:
            document = Document(io.BytesIO(raw))
        except Exception as exc:
            logger.warning("Failed to load DOCX file %s: %s", file_name, exc)
            raise FileParseError("DOCX parse failed. Please confirm file is valid or re-save as .docx.") from exc

        paragraphs = []
        for idx, paragraph in enumerate(document.paragraphs):
            self._check_timeout()
            if idx >= self._docx_max_paragraphs:
                break
            text = self._normalize_text(paragraph.text)
            if text.strip():
                paragraphs.append(text)

        table_structures: list[dict[str, Any]] = []
        table_lines: list[str] = []
        for t_idx, table in enumerate(document.tables):
            self._check_timeout()
            if t_idx >= self._docx_max_tables:
                break

            rows: list[list[str]] = []
            for r_idx, row in enumerate(table.rows):
                if r_idx > self._sample_rows:
                    break
                cells = [self._stringify_cell(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append(cells)

            if not rows:
                continue

            header = rows[0]
            sample_rows = rows[1 : self._sample_rows + 1] if len(rows) > 1 else []
            table_structures.append(
                {
                    "name": f"table_{t_idx + 1}",
                    "columns": header[:50],
                    "sample_rows": sample_rows,
                }
            )
            table_lines.extend([" | ".join(row[:10]) for row in sample_rows[:2] if any(row)])

        merged = "\n".join(paragraphs + table_lines).strip()
        if not merged:
            raise FileParseError("No readable content extracted from DOCX.")

        sample_content = paragraphs[:6] + table_lines[:4]
        return ParsedFileContent(
            file_name=Path(file_name).name,
            file_type="docx",
            mime_type=mime_type,
            size_bytes=len(raw),
            parser="python-docx",
            text_summary=self._truncate_text(merged, self._max_text_chars),
            table_structures=table_structures,
            sample_content=sample_content[:10],
            warnings=[],
            metadata={
                "paragraph_count": len(paragraphs),
                "table_count": len(table_structures),
            },
        )

    def _parse_xlsx(self, *, file_name: str, mime_type: str, raw: bytes) -> ParsedFileContent:
        try:
            from openpyxl import load_workbook
        except Exception as exc:
            logger.warning("XLSX dependency missing: %s", exc)
            raise FileParseError("XLSX parser is unavailable on server. Please contact administrator.") from exc

        try:
            workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        except Exception as exc:
            logger.warning("Failed to load XLSX file %s: %s", file_name, exc)
            raise FileParseError("XLSX parse failed. Please confirm file is valid or re-save as .xlsx.") from exc

        sheet_structures: list[dict[str, Any]] = []
        sample_content: list[str] = []
        text_parts: list[str] = []

        try:
            for s_idx, sheet in enumerate(workbook.worksheets):
                self._check_timeout()
                if s_idx >= self._xlsx_max_sheets:
                    break

                rows: list[list[str]] = []
                for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                    if r_idx > self._sample_rows:
                        break
                    values = [self._stringify_cell(v) for v in row]
                    if any(values):
                        rows.append(values)

                if not rows:
                    continue

                header = rows[0]
                sample_rows = rows[1 : self._sample_rows + 1] if len(rows) > 1 else []
                sheet_structures.append(
                    {
                        "name": sheet.title,
                        "columns": header[:50],
                        "sample_rows": sample_rows,
                    }
                )

                non_empty_cols = [x for x in header if x][:12]
                if non_empty_cols:
                    text_parts.append(f"Sheet {sheet.title} columns: {', '.join(non_empty_cols)}")
                for row in sample_rows[:2]:
                    merged = " | ".join([x for x in row if x][:10])
                    if merged:
                        sample_content.append(f"{sheet.title}: {merged}")
        finally:
            workbook.close()

        if not sheet_structures:
            raise FileParseError("No readable worksheet cells found in XLSX.")

        merged_text = "\n".join(text_parts + sample_content).strip()
        return ParsedFileContent(
            file_name=Path(file_name).name,
            file_type="xlsx",
            mime_type=mime_type,
            size_bytes=len(raw),
            parser="openpyxl",
            text_summary=self._truncate_text(merged_text, self._max_text_chars),
            table_structures=sheet_structures,
            sample_content=sample_content[:10],
            warnings=[],
            metadata={"sheet_count": len(sheet_structures)},
        )

    def parse(self, *, file_name: str, content_type: str | None, raw: bytes) -> ParsedFileContent:
        if not raw:
            raise FileParseError("Uploaded file is empty.")

        safe_name = Path(str(file_name or "uploaded_file")).name or "uploaded_file"
        mime_type = self._guess_mime(safe_name, content_type)
        kind = self._detect_kind(safe_name, mime_type)

        self._start_timer()

        if kind == "doc":
            raise FileParseError("\u6682\u4e0d\u652f\u6301 .doc\uff0c\u8bf7\u8f6c\u4e3a .docx")
        if kind == "pdf":
            return self._parse_pdf(file_name=safe_name, mime_type=mime_type, raw=raw)
        if kind == "docx":
            return self._parse_docx(file_name=safe_name, mime_type=mime_type, raw=raw)
        if kind == "xlsx":
            return self._parse_xlsx(file_name=safe_name, mime_type=mime_type, raw=raw)
        if kind == "text":
            return self._parse_text(file_name=safe_name, mime_type=mime_type, raw=raw)

        logger.info("Unsupported upload type: file=%s mime=%s", safe_name, mime_type)
        raise FileParseError("Unsupported file type. Supported: .txt/.md/.csv/.json/.jsonl/.pdf/.docx/.xlsx/.xlsm")


